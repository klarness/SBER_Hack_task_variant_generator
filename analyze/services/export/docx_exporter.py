import io
import html
import re
from typing import Any
from urllib.parse import quote
from xml.etree import ElementTree

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt, RGBColor
from latex2mathml.converter import convert as latex_to_mathml
from mathml2omml import convert as mathml_to_omml

DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def build_task_docx(task: dict[str, Any]) -> tuple[bytes, str, str]:
    document = Document()
    _configure_styles(document)

    title = _clean_text(task.get("title")) or "Варианты заданий"
    document.add_heading(title, level=1)

    _add_variants(document, task)

    output = io.BytesIO()
    document.save(output)

    filename = f"{_safe_filename(title)}-variants.docx"
    ascii_filename = _ascii_filename(filename)
    return output.getvalue(), filename, ascii_filename


def content_disposition(filename: str, ascii_filename: str) -> str:
    return f'attachment; filename="{ascii_filename}"; filename*=UTF-8\'\'{quote(filename)}'


def _configure_styles(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    for section in document.sections:
        section.top_margin = Pt(56)
        section.bottom_margin = Pt(56)
        section.left_margin = Pt(56)
        section.right_margin = Pt(56)


def _add_meta(document: Document, task: dict[str, Any]) -> None:
    rows = [
        ("Предмет", task.get("subject")),
        ("Тема", task.get("topic")),
        ("Тип", task.get("task_type")),
        ("Сложность", task.get("difficulty")),
        ("Статус", task.get("status")),
    ]
    rows = [(label, _clean_text(value)) for label, value in rows if _clean_text(value)]
    if not rows:
        return

    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value


def _add_source_items(document: Document, task: dict[str, Any]) -> None:
    items = task.get("task_items") or []
    if not isinstance(items, list) or not items:
        original_text = _clean_text(task.get("original_text"))
        if not original_text:
            return

        document.add_heading("Исходное задание", level=2)
        _add_text_block(document, original_text)
        return

    document.add_heading("Исходные пункты", level=2)
    for item in sorted(items, key=lambda value: _int(value.get("order")) if isinstance(value, dict) else 0):
        if not isinstance(item, dict):
            continue
        order = _int(item.get("order"))
        content = _clean_text(item.get("content"))
        if content:
            _add_numbered_text(document, order, content)


def _add_variants(document: Document, task: dict[str, Any]) -> None:
    variants = task.get("variants") or []
    if not isinstance(variants, list) or not variants:
        document.add_heading("Варианты", level=2)
        document.add_paragraph("Варианты еще не сформированы.")
        return

    for index, variant in enumerate(sorted(variants, key=lambda value: _int(value.get("variant_number")) if isinstance(value, dict) else 0)):
        if not isinstance(variant, dict):
            continue
        if index > 0:
            document.add_page_break()

        variant_number = _int(variant.get("variant_number")) or index + 1
        heading = document.add_heading(f"Вариант {variant_number}", level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        items = variant.get("items") or []
        if not isinstance(items, list) or not items:
            document.add_paragraph("В этом варианте нет заданий.")
            continue

        for item_index, item in enumerate(items, start=1):
            if not isinstance(item, dict):
                continue

            status = _clean_text(item.get("status")) or "ready"
            content = str(item.get("content") or "").strip()
            if status == "failed":
                _add_failed_item(document, item_index, _clean_text(item.get("error_message")))
            elif content:
                _add_numbered_content(document, item_index, content)
            else:
                _add_failed_item(document, item_index, "Пустой текст задания.")


def _add_failed_item(document: Document, index: int, error_message: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"{index}. Не удалось сгенерировать этот пункт.")
    run.bold = True
    run.font.color.rgb = RGBColor(192, 57, 43)
    if error_message:
        paragraph.add_run(" ")
        details = paragraph.add_run(error_message)
        details.italic = True
        details.font.color.rgb = RGBColor(192, 57, 43)


def _add_numbered_text(document: Document, index: int, text: str) -> None:
    lines = [line for line in text.splitlines() if line.strip()]
    if not lines:
        return

    paragraph = document.add_paragraph()
    paragraph.add_run(f"{index}. ")
    _add_rich_text(paragraph, lines[0])
    for line in lines[1:]:
        paragraph.add_run().add_break()
        _add_rich_text(paragraph, line)


def _add_numbered_content(document: Document, index: int, content: str) -> None:
    if not _looks_like_html(content):
        _add_numbered_text(document, index, _clean_text(content))
        return

    root = _parse_html_fragment(content)
    if root is None:
        _add_numbered_text(document, index, _html_to_plain_text(content))
        return

    document.add_paragraph(f"{index}.")
    _add_html_blocks(document, root)


def _looks_like_html(value: str) -> bool:
    return bool(re.search(r"<\s*(p|br|strong|b|em|i|u|s|ul|ol|li|table|tr|td|th|h2|h3)\b", value, flags=re.I))


def _parse_html_fragment(value: str) -> ElementTree.Element | None:
    try:
        wrapped = f"<root>{value}</root>"
        return ElementTree.fromstring(wrapped)
    except ElementTree.ParseError as error:
        print(f"DOCX export HTML parse error: {error}")
        return None


def _add_html_blocks(document: Document, root: ElementTree.Element) -> None:
    for child in root:
        tag = _html_tag(child)

        if tag == "table":
            _add_html_table(document, child)
            continue

        if tag in {"ul", "ol"}:
            _add_html_list(document, child, ordered=tag == "ol")
            continue

        if tag in {"p", "div", "h2", "h3"}:
            paragraph = document.add_paragraph()
            if tag in {"h2", "h3"}:
                for run in paragraph.runs:
                    run.bold = True
            _add_html_inline(paragraph, child)
            continue

        text = _html_node_plain_text(child)
        if text:
            paragraph = document.add_paragraph()
            _add_rich_text(paragraph, text)


def _add_html_list(document: Document, node: ElementTree.Element, *, ordered: bool) -> None:
    index = 1
    for child in node:
        if _html_tag(child) != "li":
            continue
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{index}. " if ordered else "• ")
        _add_html_inline(paragraph, child)
        index += 1


def _add_html_table(document: Document, node: ElementTree.Element) -> None:
    rows = [
        row
        for row in node.iter()
        if _html_tag(row) == "tr"
    ]
    if not rows:
        return

    table_cells = []
    max_cols = 0
    for row in rows:
        cells = [
            cell
            for cell in list(row)
            if _html_tag(cell) in {"td", "th"}
        ]
        if cells:
            table_cells.append(cells)
            max_cols = max(max_cols, len(cells))

    if not table_cells or max_cols == 0:
        return

    table = document.add_table(rows=len(table_cells), cols=max_cols)
    table.style = "Table Grid"

    for row_index, cells in enumerate(table_cells):
        for col_index, cell_node in enumerate(cells):
            cell = table.rows[row_index].cells[col_index]
            paragraph = cell.paragraphs[0]
            _add_html_inline(paragraph, cell_node, bold=_html_tag(cell_node) == "th")


def _add_html_inline(
    paragraph,
    node: ElementTree.Element,
    *,
    bold: bool = False,
    italic: bool = False,
    underline: bool = False,
    strike: bool = False,
) -> None:
    tag = _html_tag(node)
    bold = bold or tag in {"strong", "b"}
    italic = italic or tag in {"em", "i"}
    underline = underline or tag == "u"
    strike = strike or tag in {"s", "strike"}

    if node.text:
        _add_rich_text_with_marks(paragraph, html.unescape(node.text), bold=bold, italic=italic, underline=underline, strike=strike)

    for child in node:
        child_tag = _html_tag(child)
        if child_tag == "br":
            paragraph.add_run().add_break()
        elif child_tag in {"p", "div"} and paragraph.text:
            paragraph.add_run().add_break()
            _add_html_inline(paragraph, child, bold=bold, italic=italic, underline=underline, strike=strike)
        else:
            _add_html_inline(paragraph, child, bold=bold, italic=italic, underline=underline, strike=strike)

        if child.tail:
            _add_rich_text_with_marks(paragraph, html.unescape(child.tail), bold=bold, italic=italic, underline=underline, strike=strike)


def _add_rich_text_with_marks(
    paragraph,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    underline: bool = False,
    strike: bool = False,
) -> None:
    for segment_type, value in _split_latex_segments(text):
        if segment_type == "math":
            _add_formula_text(paragraph, value)
        else:
            run = paragraph.add_run(value)
            run.bold = bold
            run.italic = italic
            run.underline = underline
            run.font.strike = strike


def _html_tag(node: ElementTree.Element) -> str:
    return str(node.tag).lower()


def _html_node_plain_text(node: ElementTree.Element) -> str:
    parts = []
    if node.text:
        parts.append(node.text)
    for child in node:
        if _html_tag(child) == "br":
            parts.append("\n")
        else:
            parts.append(_html_node_plain_text(child))
        if child.tail:
            parts.append(child.tail)
    return _clean_text(html.unescape(" ".join(parts)))


def _html_to_plain_text(value: str) -> str:
    text = html.unescape(value)
    text = re.sub(r"(?is)<\s*br\s*/?\s*>", "\n", text)
    text = re.sub(r"(?is)<\s*/\s*(p|div|tr)\s*>", "\n", text)
    text = re.sub(r"(?is)<\s*/\s*(td|th)\s*>", " | ", text)
    text = re.sub(r"(?is)<[^>]+>", "", text)
    text = re.sub(r"\s+\|", " |", text)
    text = re.sub(r"\|\s+\|", "|", text)
    return _clean_text(text)


def _add_text_block(document: Document, text: str) -> None:
    for part in re.split(r"\n{2,}", text):
        part = part.strip()
        if part:
            paragraph = document.add_paragraph()
            _add_rich_text(paragraph, part)


def _add_rich_text(paragraph, text: str) -> None:
    for segment_type, value in _split_latex_segments(text):
        if segment_type == "math":
            _add_formula_text(paragraph, value)
        else:
            paragraph.add_run(value)


def _split_latex_segments(text: str) -> list[tuple[str, str]]:
    segments: list[tuple[str, str]] = []
    pattern = re.compile(r"\$([^$\n]+)\$")
    last_index = 0

    for match in pattern.finditer(text):
        if match.start() > last_index:
            segments.append(("text", text[last_index:match.start()]))
        segments.append(("math", match.group(1).strip()))
        last_index = match.end()

    if last_index < len(text):
        segments.append(("text", text[last_index:]))

    return segments


def _add_formula_text(paragraph, formula: str) -> None:
    formula = _normalize_latex_formula(formula)
    if not formula:
        return

    try:
        mathml = latex_to_mathml(formula)
        omml = mathml_to_omml(mathml)
        if not omml:
            raise ValueError("Empty OMML")
        paragraph._p.append(_omml_element(omml))
    except Exception:
        _add_formula_runs(paragraph, formula)


def _omml_element(omml: str):
    if omml.startswith("<m:oMath"):
        omml = omml.replace("<m:oMath", f"<m:oMath {nsdecls('m')}", 1)
    return parse_xml(omml)


def _add_formula_runs(paragraph, formula: str) -> None:
    index = 0

    while index < len(formula):
        if formula.startswith(r"\cdot", index):
            _add_math_run(paragraph, "·")
            index += len(r"\cdot")
            continue

        if formula.startswith(r"\times", index):
            _add_math_run(paragraph, "×")
            index += len(r"\times")
            continue

        if formula.startswith(r"\left", index):
            index += len(r"\left")
            continue

        if formula.startswith(r"\right", index):
            index += len(r"\right")
            continue

        if formula.startswith(r"\frac", index):
            index += len(r"\frac")
            numerator, index = _read_latex_group(formula, index)
            denominator, index = _read_latex_group(formula, index)
            _add_math_run(paragraph, "(")
            _add_formula_runs(paragraph, numerator)
            _add_math_run(paragraph, ")/(")
            _add_formula_runs(paragraph, denominator)
            _add_math_run(paragraph, ")")
            continue

        if formula.startswith(r"\sqrt", index):
            index += len(r"\sqrt")
            expression, index = _read_latex_group(formula, index)
            _add_math_run(paragraph, "√(")
            _add_formula_runs(paragraph, expression)
            _add_math_run(paragraph, ")")
            continue

        char = formula[index]
        if char == "^":
            value, index = _read_latex_group(formula, index + 1)
            _add_math_run(paragraph, _latex_to_plain(value), superscript=True)
            continue

        if char == "_":
            value, index = _read_latex_group(formula, index + 1)
            _add_math_run(paragraph, _latex_to_plain(value), subscript=True)
            continue

        if char == "\\":
            command_match = re.match(r"\\([A-Za-z]+)", formula[index:])
            if command_match:
                _add_math_run(paragraph, command_match.group(1))
                index += len(command_match.group(0))
                continue

        next_index = _next_formula_control_index(formula, index)
        _add_math_run(paragraph, formula[index:next_index])
        index = next_index


def _add_math_run(paragraph, text: str, *, superscript: bool = False, subscript: bool = False) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    run.font.name = "Cambria Math"
    run.font.superscript = superscript
    run.font.subscript = subscript


def _read_latex_group(formula: str, index: int) -> tuple[str, int]:
    while index < len(formula) and formula[index].isspace():
        index += 1

    if index >= len(formula):
        return "", index

    if formula[index] != "{":
        return formula[index], index + 1

    depth = 0
    start = index + 1
    for position in range(index, len(formula)):
        if formula[position] == "{":
            depth += 1
        elif formula[position] == "}":
            depth -= 1
            if depth == 0:
                return formula[start:position], position + 1

    return formula[start:], len(formula)


def _next_formula_control_index(formula: str, index: int) -> int:
    controls = [position for position in (formula.find(symbol, index + 1) for symbol in ["\\", "^", "_"]) if position != -1]
    return min(controls) if controls else len(formula)


def _normalize_latex_formula(formula: str) -> str:
    formula = formula.replace(r"\,", " ")
    formula = formula.replace(r"\;", " ")
    formula = formula.replace(r"\:", " ")
    formula = formula.replace(r"\!", "")
    formula = re.sub(r"\s+", " ", formula)
    return formula.strip()


def _latex_to_plain(value: str) -> str:
    value = _normalize_latex_formula(value)
    value = value.replace(r"\cdot", "·").replace(r"\times", "×")
    value = value.replace("{", "").replace("}", "")
    value = value.replace("\\", "")
    return value


def _clean_text(value: Any) -> str:
    if value is None:
        return ""
    text = str(value)
    text = text.replace("\x00", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _safe_filename(value: str) -> str:
    value = _clean_text(value)
    value = re.sub(r'[<>:"/\\|?*\x00-\x1F]+', "-", value)
    value = re.sub(r"\s+", " ", value).strip(" .-_")
    return (value[:80].strip(" .-_") or "task")


def _ascii_filename(filename: str) -> str:
    stem, dot, suffix = filename.rpartition(".")
    ascii_stem = stem.encode("ascii", errors="ignore").decode()
    ascii_stem = re.sub(r"[^A-Za-z0-9._-]+", "-", ascii_stem).strip("-._")
    if not ascii_stem:
        ascii_stem = "task-variants"
    if dot and suffix:
        return f"{ascii_stem}.{suffix}"
    return ascii_stem


def _int(value: Any) -> int:
    try:
        return int(value)
    except (TypeError, ValueError):
        return 0
