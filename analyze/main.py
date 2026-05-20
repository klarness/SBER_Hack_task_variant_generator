import io
import json
import os
from typing import Any

from docx import Document
from fastapi import Body, FastAPI, File, Form, HTTPException, Response, UploadFile

from analyze.schemas.response import (
    AnalyzeItem,
    AnalyzeResponse,
    GenerateRequest,
    GenerateResponse,
    ParseResponse,
    ValidateRequest,
    ValidateResponse,
)
from analyze.services.llm.client import GigaChatClient
from analyze.services.parsing.extraction import FileExtractionService
from analyze.services.parsing.normalizer import TextNormalizer

app = FastAPI(title="Analyze Service")

MAX_PARSE_FILE_SIZE = int(os.getenv("MAX_PARSE_FILE_SIZE", str(32 * 1024 * 1024)))

extraction_service = FileExtractionService()
llm_client = GigaChatClient()


@app.get("/healthz")
async def healthz():
    return {"status": "ok"}


@app.post("/parse", response_model=ParseResponse)
async def parse(file: UploadFile = File(...)):
    await file.seek(0, io.SEEK_END)
    size = await file.tell()
    await file.seek(0)

    if size > MAX_PARSE_FILE_SIZE:
        raise HTTPException(
            status_code=413,
            detail=f"File too large: {size} bytes, max allowed is {MAX_PARSE_FILE_SIZE} bytes",
        )

    raw_text = await extraction_service.parse(file)
    normalized_text = TextNormalizer.normalize(raw_text)
    return ParseResponse(raw_text=raw_text, normalized_text=normalized_text)


@app.post("/analyze", response_model=AnalyzeResponse)
async def analyze(
    files: list[UploadFile] | None = File(default=None),
    text: str = Form(default=""),
    title: str = Form(default=""),
    settings: str = Form(default="{}"),
    user_id: str = Form(default=""),
):
    parsed_parts: list[str] = []

    for file in files or []:
        raw_text = await extraction_service.parse(file)
        normalized_text = TextNormalizer.normalize(raw_text)
        if normalized_text:
            parsed_parts.append(normalized_text)

    if text.strip():
        parsed_parts.append(TextNormalizer.normalize(text))

    original_text = "\n\n".join(parsed_parts).strip()
    if not original_text:
        raise HTTPException(status_code=400, detail="No text was extracted from request")

    try:
        llm_result = await llm_client.analyze_task(
            original_text=original_text,
            title=title,
            settings=_parse_settings(settings),
        )
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"GigaChat analyze failed: {exc}") from exc

    return AnalyzeResponse(
        original_text=original_text,
        items=_parse_llm_items(llm_result),
        subject=str(llm_result.get("subject") or ""),
        topic=str(llm_result.get("topic") or ""),
        task_type=str(llm_result.get("task_type") or ""),
        difficulty=str(llm_result.get("difficulty") or ""),
    )


@app.post("/generate", response_model=GenerateResponse)
async def generate(request: GenerateRequest):
    if not request.source_content.strip():
        raise HTTPException(status_code=400, detail="source_content is required")

    try:
        content = await llm_client.generate_variant(request.model_dump())
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"GigaChat generate failed: {exc}") from exc

    return GenerateResponse(content=content)


@app.post("/validate", response_model=ValidateResponse)
async def validate(request: ValidateRequest):
    try:
        valid = await llm_client.validate_variant(request.model_dump())
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"GigaChat validate failed: {exc}") from exc

    return ValidateResponse(valid=valid)


@app.post("/export")
async def export(task: dict[str, Any] = Body(...)):
    document = Document()
    title = task.get("title") or "Task variants"
    document.add_heading(title, level=1)

    original_text = task.get("original_text") or ""
    if original_text:
        document.add_heading("Original task", level=2)
        document.add_paragraph(original_text)

    for variant in task.get("variants") or []:
        variant_number = variant.get("variant_number", "")
        document.add_heading(f"Variant {variant_number}", level=2)
        for index, item in enumerate(variant.get("items") or [], start=1):
            document.add_paragraph(f"{index}. {item.get('content', '')}")

    output = io.BytesIO()
    document.save(output)
    filename = f"{title[:40].strip() or 'task'}-variants.docx"

    return Response(
        content=output.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "X-Filename": filename,
            "Content-Disposition": f'attachment; filename="{filename}"',
        },
    )


def _parse_settings(settings: str) -> dict[str, Any]:
    try:
        parsed = json.loads(settings or "{}")
        if isinstance(parsed, dict):
            return parsed
    except json.JSONDecodeError:
        pass
    return {}


def _parse_llm_items(llm_result: dict[str, Any]) -> list[AnalyzeItem]:
    raw_items = llm_result.get("items")
    if not isinstance(raw_items, list) or not raw_items:
        raise HTTPException(status_code=502, detail="GigaChat analyze response does not contain items")

    items: list[AnalyzeItem] = []
    for index, raw_item in enumerate(raw_items, start=1):
        if not isinstance(raw_item, dict):
            raise HTTPException(status_code=502, detail="GigaChat analyze item must be an object")

        content = str(raw_item.get("content") or "").strip()
        if not content:
            raise HTTPException(status_code=502, detail="GigaChat analyze item content is empty")

        try:
            order = int(raw_item.get("order"))
        except (TypeError, ValueError):
            order = index

        items.append(
            AnalyzeItem(
                order=order,
                context=str(raw_item.get("context") or ""),
                content=content,
            )
        )

    items.sort(key=lambda item: item.order)
    return items
